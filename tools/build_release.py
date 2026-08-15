"""Build release HTML, DOCX, and PDF from canonical Markdown report."""
from __future__ import annotations

import argparse
import html
import os
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "report" / "REPORT_FINAL.md"
HTML_OUT = ROOT / "report" / "REPORT_FINAL.html"
DOCX_OUT = ROOT / "report" / "BaoCao_TieuLuan_PLDC_UIT_FINAL.docx"
PDF_OUT = ROOT / "report" / "BaoCao_TieuLuan_PLDC_UIT_FINAL.pdf"


# ==============================================================================
# HTML BUILDER
# ==============================================================================

CSS = """
@page { size: A4; margin: 2.2cm 2.0cm 2.2cm 3.2cm; }
@media print {
    body { font-size: 13pt; line-height: 1.5; }
    .page-break { page-break-before: always; }
    .no-print { display: none; }
}
body {
    font-family: 'Times New Roman', Times, serif;
    font-size: 13pt;
    line-height: 1.5;
    color: #111;
    max-width: 210mm;
    margin: 0 auto;
    padding: 2.2cm 2.0cm 2.2cm 3.2cm;
    background: #fff;
    box-sizing: border-box;
}
.cover-page {
    text-align: center;
    min-height: 24cm;
    display: flex;
    flex-direction: column;
    justify-content: space-between;
    page-break-after: always;
}
.cover-title { font-size: 16pt; font-weight: bold; text-transform: uppercase; margin-top: 1.5cm; line-height: 1.4; }
.cover-subtitle { font-size: 14pt; font-weight: bold; margin: 0.8cm 0 1.5cm; line-height: 1.4; }
.cover-meta { font-size: 12pt; text-align: left; margin: 2cm 0; line-height: 1.6; }
.cover-footer { font-size: 12pt; text-align: center; font-weight: bold; margin-bottom: 1cm; }
h1 {
    text-align: center;
    font-size: 15pt;
    font-weight: bold;
    text-transform: uppercase;
    margin: 24pt 0 14pt;
    page-break-before: always;
    page-break-after: avoid;
}
h2 {
    font-size: 13.5pt;
    font-weight: bold;
    text-align: left;
    margin: 16pt 0 8pt;
    page-break-after: avoid;
}
h3 {
    font-size: 13pt;
    font-weight: bold;
    text-align: left;
    margin: 12pt 0 6pt;
    page-break-after: avoid;
}
h4 {
    font-size: 13pt;
    font-style: italic;
    font-weight: bold;
    text-align: left;
    margin: 10pt 0 4pt;
    page-break-after: avoid;
}
p {
    text-align: justify;
    text-indent: 1.25cm;
    margin: 0 0 6pt 0;
}
.no-indent {
    text-indent: 0 !important;
}
ul, ol {
    margin: 4pt 0 8pt 0;
    padding-left: 2cm;
}
li {
    text-align: justify;
    margin-bottom: 4pt;
}
table {
    width: 100%;
    border-collapse: collapse;
    font-size: 10pt;
    margin: 12pt 0;
    page-break-inside: avoid;
}
th, td {
    border: 0.6pt solid #333;
    padding: 5pt 6pt;
    vertical-align: top;
}
th {
    background: #ececec;
    text-align: center;
    font-weight: bold;
}
td {
    text-align: left;
}
pre {
    white-space: pre-wrap;
    font-family: 'Courier New', monospace;
    font-size: 9pt;
    line-height: 1.25;
    padding: 8pt;
    background: #fdfdfd;
    border: 0.6pt solid #999;
    margin: 10pt 0;
}
hr {
    border: 0;
    border-top: 0.7pt solid #555;
    margin: 16pt 0;
}
.table-caption {
    font-weight: bold;
    font-size: 11pt;
    text-align: left;
    margin: 10pt 0 4pt 0;
    text-indent: 0;
}
.toc-box {
    margin: 20pt 0;
    padding: 10pt 0;
    page-break-after: always;
}
.toc-item {
    display: flex;
    justify-content: space-between;
    margin: 4pt 0;
    text-indent: 0;
}
.toc-l1 { font-weight: bold; margin-top: 8pt; }
.toc-l2 { padding-left: 1.5cm; }
.toc-l3 { padding-left: 2.5cm; font-size: 11pt; }
.dots { flex-grow: 1; border-bottom: 1px dotted #888; margin: 0 6pt 4pt; }
"""


def inline_html(value: str) -> str:
    value = html.escape(value, quote=False)
    value = re.sub(r"`([^`]+)`", r"<code>\1</code>", value)
    value = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", value)
    value = re.sub(r"\*([^*]+)\*", r"<em>\1</em>", value)
    return value


def generate_toc_html(text: str) -> str:
    """Generate a clean hierarchical Table of Contents for HTML."""
    toc_lines = ['<div class="toc-box">', '<h1>MỤC LỤC</h1>']
    lines = text.splitlines()
    in_cover = True
    for line in lines:
        if line.startswith("# MỞ ĐẦU"):
            in_cover = False
        if in_cover:
            continue
        m = re.match(r"^(#{1,2})\s+(.*)$", line)
        if m:
            lvl = len(m.group(1))
            title = m.group(2).strip()
            clean_title = re.sub(r"[*_`]", "", title)
            if lvl == 1:
                toc_lines.append(f'<div class="toc-item toc-l1"><span>{clean_title}</span><span class="dots"></span></div>')
            elif lvl == 2:
                toc_lines.append(f'<div class="toc-item toc-l2"><span>{clean_title}</span><span class="dots"></span></div>')
    toc_lines.append('</div>')
    return "\n".join(toc_lines)


def build_html(text: str) -> str:
    lines = text.splitlines()
    out = []
    in_pre = False
    in_table = False
    list_kind = None
    in_cover = True
    cover_lines = []

    def close_list():
        nonlocal list_kind
        if list_kind:
            out.append(f"</{list_kind}>")
            list_kind = None

    def close_table():
        nonlocal in_table
        if in_table:
            out.append("</tbody></table>")
            in_table = False

    # First collect cover info and body
    body_lines = []
    for line in lines:
        if line.startswith("# MỞ ĐẦU"):
            in_cover = False
        if in_cover:
            cover_lines.append(line)
        else:
            body_lines.append(line)

    # Process cover & member table
    out.append('<div class="cover-page">')
    out.append('<div>')
    out.append('<div style="font-size:13pt;font-weight:bold;margin-bottom:4pt;">ĐẠI HỌC QUỐC GIA THÀNH PHỐ HỒ CHÍ MINH</div>')
    out.append('<div style="font-size:14pt;font-weight:bold;margin-bottom:12pt;">TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN</div>')
    out.append('<div style="font-size:13pt;font-weight:bold;margin-bottom:40pt;">KHOA / BỘ MÔN LÝ LUẬN CHÍNH TRỊ</div>')
    out.append('<div class="cover-title">BÁO CÁO TIỂU LUẬN MÔN HỌC<br>PHÁP LUẬT ĐẠI CƯƠNG</div>')
    out.append('<div class="cover-subtitle">ĐỀ TÀI:<br>PHÂN TÍCH CÁC YẾU TỐ CẤU THÀNH VI PHẠM PHÁP LUẬT TỪ VỤ VIỆC QUẢNG CÁO GIAN DỐI SẢN PHẨM ĐÔNG Y TRÊN MẠNG XÃ HỘI</div>')
    out.append('</div>')
    out.append('<div class="cover-meta">')
    out.append('<p class="no-indent"><strong>Giảng viên hướng dẫn:</strong> [PLACEHOLDER]</p>')
    out.append('<p class="no-indent"><strong>Lớp học phần:</strong> [PLACEHOLDER]</p>')
    out.append('<p class="no-indent"><strong>Nhóm sinh viên thực hiện:</strong> [PLACEHOLDER]</p>')
    out.append('</div>')
    out.append('<div class="cover-footer">THÀNH PHỐ HỒ CHÍ MINH, NĂM 2026</div>')
    out.append('</div>')

    # Add Member table
    out.append('<div style="page-break-after:always;padding-top:1cm;">')
    out.append('<h2 style="text-align:center;">DANH SÁCH THÀNH VIÊN VÀ PHÂN CÔNG NHIỆM VỤ</h2>')
    out.append('<table><thead><tr><th>STT</th><th>Họ và tên</th><th>MSSV</th><th>Vai trò</th><th>Nội dung phụ trách</th></tr></thead><tbody>')
    out.append('<tr><td style="text-align:center;">1</td><td>[PLACEHOLDER]</td><td>[PLACEHOLDER]</td><td>Nhóm trưởng</td><td>Điều phối và tổng hợp báo cáo</td></tr>')
    out.append('<tr><td style="text-align:center;">2</td><td>[PLACEHOLDER]</td><td>[PLACEHOLDER]</td><td>Thành viên</td><td>Cơ sở lý luận và khách thể</td></tr>')
    out.append('<tr><td style="text-align:center;">3</td><td>[PLACEHOLDER]</td><td>[PLACEHOLDER]</td><td>Thành viên</td><td>Khái quát vụ việc và dòng thời gian</td></tr>')
    out.append('<tr><td style="text-align:center;">4</td><td>[PLACEHOLDER]</td><td>[PLACEHOLDER]</td><td>Thành viên</td><td>Chủ thể và mặt chủ quan</td></tr>')
    out.append('<tr><td style="text-align:center;">5</td><td>[PLACEHOLDER]</td><td>[PLACEHOLDER]</td><td>Thành viên</td><td>Mặt khách quan, kết luận và tài liệu tham khảo</td></tr>')
    out.append('</tbody></table>')
    out.append('</div>')

    # Add Table of Contents
    out.append(generate_toc_html(text))

    # Process body lines
    for i, line in enumerate(body_lines):
        if line.startswith("```"):
            close_list()
            close_table()
            out.append("</pre>" if in_pre else "<pre>")
            in_pre = not in_pre
            continue
        if in_pre:
            out.append(html.escape(line))
            continue
        if re.match(r"^\|.*\|\s*$", line):
            close_list()
            parts = [x.strip() for x in line.strip()[1:-1].split("|")]
            if i + 1 < len(body_lines) and re.match(r"^\|\s*:?-+", body_lines[i + 1]):
                out.append('<table><thead><tr>' + "".join(f"<th>{inline_html(x)}</th>" for x in parts) + "</tr></thead><tbody>")
                in_table = True
            elif not re.match(r"^\|\s*:?-+", line):
                out.append("<tr>" + "".join(f"<td>{inline_html(x)}</td>" for x in parts) + "</tr>")
            continue
        close_table()
        if not line.strip():
            close_list()
            continue
        if line.strip() == "---":
            close_list()
            out.append("<hr>")
            continue

        # Check Table Caption
        m_table = re.match(r"^\*\*(Bảng \d+\..*?)\*\*$", line)
        if m_table:
            close_list()
            out.append(f'<p class="table-caption">{inline_html(m_table.group(1))}</p>')
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            close_list()
            n = len(m.group(1))
            out.append(f"<h{n}>{inline_html(m.group(2))}</h{n}>")
            continue

        # Lists
        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            if list_kind != "ul":
                close_list()
                out.append("<ul>")
                list_kind = "ul"
            out.append(f"<li>{inline_html(m.group(1))}</li>")
            continue
        m = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if m:
            if list_kind != "ol":
                close_list()
                out.append("<ol>")
                list_kind = "ol"
            out.append(f"<li>{inline_html(m.group(1))}</li>")
            continue

        # Blockquote / Note
        if line.startswith(">"):
            close_list()
            out.append(f'<p style="font-style:italic;margin:6pt 0;padding-left:1cm;">{inline_html(line[1:].strip())}</p>')
            continue

        close_list()
        out.append(f"<p>{inline_html(line)}</p>")

    close_list()
    close_table()

    return f"""<!doctype html>
<html lang="vi">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Tiểu luận Pháp luật đại cương - UIT</title>
<style>{CSS}</style>
</head>
<body>
{"\n".join(out)}
</body>
</html>"""


# ==============================================================================
# DOCX BUILDER
# ==============================================================================

def set_run_font(run, name="Times New Roman", size=13, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = color


def add_formatted_runs(paragraph, text: str, default_size=13):
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    for piece in pattern.split(text):
        if not piece:
            continue
        bold = piece.startswith("**") and piece.endswith("**")
        italic = piece.startswith("*") and piece.endswith("*") and not bold
        code = piece.startswith("`") and piece.endswith("`")
        content = piece[2:-2] if bold else piece[1:-1] if (italic or code) else piece
        run = paragraph.add_run(content)
        set_run_font(run, name="Courier New" if code else "Times New Roman", size=default_size - 1 if code else default_size, bold=bold, italic=italic)


def set_cell_properties(cell, text: str, is_header=False, font_size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
    add_formatted_runs(p, text, default_size=font_size)
    if is_header:
        for r in p.runs:
            r.bold = True
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_native_toc_field(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    fldSimple = OxmlElement("w:fldSimple")
    fldSimple.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    paragraph._p.append(fldSimple)


def set_table_col_widths(table, col_widths):
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            if idx < len(row.cells):
                row.cells[idx].width = width


def build_docx(text: str):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3.2)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)

    # Normal Style
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Heading Styles
    heading_configs = [
        (1, 15, True, False, WD_ALIGN_PARAGRAPH.CENTER, 18, 10, True),
        (2, 13.5, True, False, WD_ALIGN_PARAGRAPH.LEFT, 14, 6, False),
        (3, 13, True, False, WD_ALIGN_PARAGRAPH.LEFT, 10, 4, False),
        (4, 13, True, True, WD_ALIGN_PARAGRAPH.LEFT, 8, 3, False),
    ]
    for lvl, size, bold, italic, align, sp_before, sp_after, page_break in heading_configs:
        st = doc.styles[f"Heading {lvl}"]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.italic = italic
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(sp_before)
        st.paragraph_format.space_after = Pt(sp_after)
        st.paragraph_format.line_spacing = 1.3
        st.paragraph_format.alignment = align
        st.paragraph_format.first_line_indent = Cm(0)
        st.paragraph_format.keep_with_next = True
        if page_break and lvl == 1:
            st.paragraph_format.page_break_before = True

    # Header / Footer (No page number on cover)
    sec.different_first_page_header_footer = True
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.first_line_indent = Cm(0)
    r_ft = footer.add_run("Trang ")
    set_run_font(r_ft, size=10)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)

    # 1. BUILD COVER PAGE
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.first_line_indent = Cm(0)
    p1.paragraph_format.space_after = Pt(3)
    r = p1.add_run("ĐẠI HỌC QUỐC GIA THÀNH PHỐ HỒ CHÍ MINH\nTRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN")
    set_run_font(r, size=14, bold=True)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.first_line_indent = Cm(0)
    p_sub.paragraph_format.space_after = Pt(36)
    r_sub = p_sub.add_run("KHOA / BỘ MÔN LÝ LUẬN CHÍNH TRỊ")
    set_run_font(r_sub, size=13, bold=True)

    p_tt = doc.add_paragraph()
    p_tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tt.paragraph_format.first_line_indent = Cm(0)
    p_tt.paragraph_format.space_after = Pt(14)
    r_tt = p_tt.add_run("BÁO CÁO TIỂU LUẬN MÔN HỌC\nPHÁP LUẬT ĐẠI CƯƠNG")
    set_run_font(r_tt, size=16, bold=True)

    p_de = doc.add_paragraph()
    p_de.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_de.paragraph_format.first_line_indent = Cm(0)
    p_de.paragraph_format.space_after = Pt(72)
    r_de = p_de.add_run("ĐỀ TÀI:\nPHÂN TÍCH CÁC YẾU TỐ CẤU THÀNH VI PHẠM PHÁP LUẬT TỪ VỤ VIỆC QUẢNG CÁO GIAN DỐI SẢN PHẨM ĐÔNG Y TRÊN MẠNG XÃ HỘI")
    set_run_font(r_de, size=14, bold=True)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.first_line_indent = Cm(0)
    p_meta.paragraph_format.left_indent = Cm(1.5)
    p_meta.paragraph_format.line_spacing = 1.4
    p_meta.paragraph_format.space_after = Pt(60)
    r_meta = p_meta.add_run(
        "Giảng viên hướng dẫn:  [PLACEHOLDER]\n"
        "Lớp học phần:          [PLACEHOLDER]\n"
        "Nhóm sinh viên:        [PLACEHOLDER]\n"
        "Sinh viên thực hiện:   [PLACEHOLDER]"
    )
    set_run_font(r_meta, size=12.5)

    p_bot = doc.add_paragraph()
    p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bot.paragraph_format.first_line_indent = Cm(0)
    p_bot.paragraph_format.space_after = Pt(0)
    r_bot = p_bot.add_run("THÀNH PHỐ HỒ CHÍ MINH, NĂM 2026")
    set_run_font(r_bot, size=12, bold=True)

    doc.add_page_break()

    # 2. MEMBER TABLE PAGE
    p_mhead = doc.add_paragraph()
    p_mhead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_mhead.paragraph_format.first_line_indent = Cm(0)
    p_mhead.paragraph_format.space_before = Pt(12)
    p_mhead.paragraph_format.space_after = Pt(12)
    r_mhead = p_mhead.add_run("DANH SÁCH THÀNH VIÊN VÀ PHÂN CÔNG NHIỆM VỤ")
    set_run_font(r_mhead, size=14, bold=True)

    member_data = [
        ["STT", "Họ và tên", "MSSV", "Vai trò", "Nội dung phụ trách"],
        ["1", "[PLACEHOLDER]", "[PLACEHOLDER]", "Nhóm trưởng", "Điều phối và tổng hợp báo cáo"],
        ["2", "[PLACEHOLDER]", "[PLACEHOLDER]", "Thành viên", "Cơ sở lý luận và khách thể"],
        ["3", "[PLACEHOLDER]", "[PLACEHOLDER]", "Thành viên", "Khái quát vụ việc và dòng thời gian"],
        ["4", "[PLACEHOLDER]", "[PLACEHOLDER]", "Thành viên", "Chủ thể và mặt chủ quan"],
        ["5", "[PLACEHOLDER]", "[PLACEHOLDER]", "Thành viên", "Mặt khách quan, kết luận và tài liệu tham khảo"],
    ]
    t_m = doc.add_table(rows=len(member_data), cols=5)
    t_m.style = "Table Grid"
    t_m.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(member_data):
        for c_idx, val in enumerate(row):
            set_cell_properties(t_m.cell(r_idx, c_idx), val, is_header=(r_idx == 0), font_size=10)
    set_table_col_widths(t_m, [Cm(1.2), Cm(3.8), Cm(2.4), Cm(2.8), Cm(5.6)])

    doc.add_page_break()

    # 3. WORD TABLE OF CONTENTS (Not using Heading 1 so it does not index itself)
    p_toc_head = doc.add_paragraph()
    p_toc_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_toc_head.paragraph_format.first_line_indent = Cm(0)
    p_toc_head.paragraph_format.space_before = Pt(18)
    p_toc_head.paragraph_format.space_after = Pt(14)
    r_toc_head = p_toc_head.add_run("MỤC LỤC")
    set_run_font(r_toc_head, size=15, bold=True)

    p_toc_field = doc.add_paragraph()
    p_toc_field.paragraph_format.first_line_indent = Cm(0)
    p_toc_field.paragraph_format.line_spacing = 1.3
    p_toc_field.paragraph_format.space_after = Pt(12)
    add_native_toc_field(p_toc_field)

    doc.add_page_break()

    # 4. PARSE CANONICAL BODY
    lines = text.splitlines()
    in_body = False
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# MỞ ĐẦU"):
            in_body = True
        if not in_body:
            i += 1
            continue

        # Code block
        if line.startswith("```"):
            i += 1
            code_block = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_block.append(lines[i])
                i += 1
            p_code = doc.add_paragraph()
            p_code.paragraph_format.left_indent = Cm(0.5)
            p_code.paragraph_format.first_line_indent = Cm(0)
            p_code.paragraph_format.line_spacing = 1.0
            p_code.paragraph_format.space_before = Pt(4)
            p_code.paragraph_format.space_after = Pt(6)
            r_c = p_code.add_run("\n".join(code_block))
            set_run_font(r_c, name="Courier New", size=8.5)
            i += 1
            continue

        # Markdown Table
        if line.startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].startswith("|"):
                if not re.match(r"^\|\s*:?-+", lines[i]):
                    table_rows.append([x.strip() for x in lines[i].strip()[1:-1].split("|")])
                i += 1
            if table_rows:
                tbl = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                tbl.style = "Table Grid"
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                num_cols = len(table_rows[0])
                for r_idx, row in enumerate(table_rows):
                    for c_idx, val in enumerate(row):
                        set_cell_properties(tbl.cell(r_idx, c_idx), val, is_header=(r_idx == 0), font_size=9.5)
                # Auto balance column widths
                total_w = Cm(15.8)
                if num_cols == 3:
                    set_table_col_widths(tbl, [Cm(3.8), Cm(7.2), Cm(4.8)])
                elif num_cols == 4:
                    set_table_col_widths(tbl, [Cm(3.2), Cm(4.6), Cm(4.0), Cm(4.0)])
                elif num_cols == 5:
                    set_table_col_widths(tbl, [Cm(1.2), Cm(3.8), Cm(2.4), Cm(2.8), Cm(5.6)])
                else:
                    col_w = total_w / num_cols
                    set_table_col_widths(tbl, [col_w] * num_cols)
                for r_idx, row in enumerate(tbl.rows):
                    trPr = row._tr.get_or_add_trPr()
                    trPr.append(OxmlElement('w:cantSplit'))
                    if r_idx == 0:
                        trPr.append(OxmlElement('w:tblHeader'))
            continue

        # Blank or HR
        if not line.strip() or line.strip() == "---":
            i += 1
            continue

        # Table caption
        m_table = re.match(r"^\*\*(Bảng \d+\..*?)\*\*$", line)
        if m_table:
            p_tc = doc.add_paragraph()
            p_tc.paragraph_format.first_line_indent = Cm(0)
            p_tc.paragraph_format.space_before = Pt(10)
            p_tc.paragraph_format.space_after = Pt(3)
            p_tc.paragraph_format.keep_with_next = True
            r_tc = p_tc.add_run(m_table.group(1))
            set_run_font(r_tc, size=11, bold=True)
            i += 1
            continue

        # Headings
        m_h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m_h:
            lvl = len(m_h.group(1))
            title = m_h.group(2).strip()
            target_lvl = min(lvl, 4)
            p_h = doc.add_paragraph(style=f"Heading {target_lvl}")
            if target_lvl == 1:
                p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_formatted_runs(p_h, title, default_size=[15, 13.5, 13, 13][target_lvl - 1])
            i += 1
            continue

        # Unordered list
        m_ul = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m_ul:
            p_l = doc.add_paragraph(style="List Bullet")
            p_l.paragraph_format.first_line_indent = Cm(0)
            p_l.paragraph_format.left_indent = Cm(1.25)
            p_l.paragraph_format.space_after = Pt(3)
            p_l.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_formatted_runs(p_l, m_ul.group(1), default_size=13)
            i += 1
            continue

        # Ordered list
        m_ol = re.match(r"^\s*(\d+)[.)]\s+(.*)$", line)
        if m_ol:
            p_l = doc.add_paragraph()
            p_l.paragraph_format.first_line_indent = Cm(0)
            p_l.paragraph_format.left_indent = Cm(1.25)
            p_l.paragraph_format.space_after = Pt(3)
            p_l.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_l.add_run(f"{m_ol.group(1)}. ")
            add_formatted_runs(p_l, m_ol.group(2), default_size=13)
            i += 1
            continue

        # Blockquote
        if line.startswith(">"):
            p_q = doc.add_paragraph()
            p_q.paragraph_format.first_line_indent = Cm(0)
            p_q.paragraph_format.left_indent = Cm(1.0)
            p_q.paragraph_format.space_before = Pt(4)
            p_q.paragraph_format.space_after = Pt(6)
            p_q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_formatted_runs(p_q, line[1:].strip(), default_size=12)
            for r_q in p_q.runs:
                r_q.italic = True
            i += 1
            continue

        # Standard Paragraph
        p_body = doc.add_paragraph()
        p_body.paragraph_format.first_line_indent = Cm(1.25)
        p_body.paragraph_format.space_after = Pt(6)
        p_body.paragraph_format.line_spacing = 1.5
        p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_formatted_runs(p_body, line, default_size=13)
        i += 1

    doc.save(DOCX_OUT)
    print(f"[DOCX] Generated: {DOCX_OUT}")


# ==============================================================================
# PDF EXPORT VIA WORD AUTOMATION
# ==============================================================================

def export_docx_to_pdf_via_word():
    """Use Microsoft Word COM to update fields (TOC page numbers) and export to PDF."""
    ps_script = f"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0

try {{
    $docxPath = "{DOCX_OUT}"
    $pdfPath = "{PDF_OUT}"
    
    $doc = $word.Documents.Open($docxPath)
    
    # Update all fields in document (including Table of Contents)
    $doc.Fields.Update()
    foreach ($story in $doc.StoryRanges) {{
        $story.Fields.Update()
    }}
    
    # Save updated DOCX with real TOC
    $doc.Save()
    
    # Export to PDF (wdExportFormatPDF = 17)
    $doc.ExportAsFixedFormat($pdfPath, 17, $false, 0, 0, 1, 1, 0, $true, $true, 0, $true, $true, $false)
    
    $doc.Close($false)
    Write-Output "SUCCESS: Updated Word TOC fields and exported PDF"
}} catch {{
    Write-Output "ERROR: $($_.Exception.Message)"
}} finally {{
    $word.Quit()
    [System.Runtime.InteropServices.Marshal]::ReleaseComObject($word) | Out-Null
    [System.GC]::Collect()
    [System.GC]::WaitForPendingFinalizers()
}}
"""
    tmp_ps = ROOT / "tools" / "update_and_export.ps1"
    tmp_ps.write_text(ps_script, encoding="utf-8")
    try:
        res = subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-File", str(tmp_ps)],
            capture_output=True,
            text=True,
            timeout=45,
        )
        print(res.stdout)
        if "SUCCESS" in res.stdout:
            print(f"[PDF] Authoritative PDF exported from DOCX: {PDF_OUT}")
            return True
        else:
            print(f"[PDF] Word automation error: {res.stderr}")
            return False
    except Exception as e:
        print(f"[PDF] Word automation exception: {e}")
        return False
    finally:
        if tmp_ps.exists():
            tmp_ps.unlink()


def main():
    text = SOURCE.read_text(encoding="utf-8")

    # 1. Build HTML
    html_content = build_html(text)
    HTML_OUT.write_text(html_content, encoding="utf-8")
    print(f"[HTML] Generated: {HTML_OUT}")

    # 2. Build DOCX
    build_docx(text)

    # 3. Export PDF via Word Automation
    success = export_docx_to_pdf_via_word()
    if not success:
        print("[WARNING] Word COM export failed or timed out. Checking PDF status...")


if __name__ == "__main__":
    main()
